# -*- coding: utf-8 -*-
import markdown
from docx import Document
from bs4 import BeautifulSoup

INPUT_FILE = "sample.md"
OUTPUT_FILE = "sample.docx"

with open(INPUT_FILE, "r", encoding="utf-8") as f:
    md_text = f.read()

html = markdown.markdown(md_text, extensions=["tables"])

soup = BeautifulSoup(html, "html.parser")

doc = Document()

def handle_table(table_tag):
    rows = table_tag.find_all("tr")
    if not rows:
        return

    cols = rows[0].find_all(["th", "td"])
    table = doc.add_table(rows=1, cols=len(cols))

    # Header
    hdr_cells = table.rows[0].cells
    for i, cell in enumerate(cols):
        hdr_cells[i].text = cell.get_text(strip=True)

    # Body
    for row in rows[1:]:
        row_cells = table.add_row().cells
        cells = row.find_all(["td", "th"])
        for i, cell in enumerate(cells):
            row_cells[i].text = cell.get_text(strip=True)

for element in soup.children:
    if element.name == "h1":
        doc.add_heading(element.text, level=1)

    elif element.name == "h2":
        doc.add_heading(element.text, level=2)

    elif element.name == "h3":
        doc.add_heading(element.text, level=3)

    elif element.name == "p":
        doc.add_paragraph(element.text)

    elif element.name == "ul":
        for li in element.find_all("li"):
            doc.add_paragraph(li.text, style="List Bullet")

    elif element.name == "ol":
        for li in element.find_all("li"):
            doc.add_paragraph(li.text, style="List Number")

    elif element.name == "table":
        handle_table(element)

doc.save(OUTPUT_FILE)

print(f"Success {OUTPUT_FILE}")